from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

IMAGES_DIR = os.path.join(os.path.dirname(__file__), 'docs', 'images')
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), 'docs', 'Manual_de_Usuario_Mercurio_Software.docx')

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(7, 94, 84)  # WhatsApp green theme
    return h

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_image(doc, filename, caption='', width=Inches(5.5)):
    path = os.path.join(IMAGES_DIR, filename)
    if not os.path.exists(path):
        add_body(doc, f'[Imagen no encontrada: {filename}]')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.style.font.size = Pt(9)
        cap.style.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph()  # spacer

def add_bullet(doc, text, bold_prefix=''):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def build_document():
    doc = Document()

    # ====== STYLES ======
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # ====== PORTADA ======
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Manual de Usuario')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(7, 94, 84)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Mercurio Software\nPlataforma de Envío Masivo de Mensajes WhatsApp')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    doc.add_paragraph()

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run('Versión 1.0\nJulio 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_page_break()

    # ====== TABLA DE CONTENIDO ======
    add_heading(doc, 'Tabla de Contenido', 1)
    toc_items = [
        '1. Introducción',
        '2. Requisitos del Sistema',
        '3. Inicio de Sesión',
        '4. Pantalla Principal',
        '5. Gestión de Plantillas',
        '6. Configuración de Mensajes',
        '7. Prospectos y Envío Masivo',
        '8. Chat en Tiempo Real',
        '9. Cargue de Archivos',
        '10. Historial de Mensajes',
        '11. Soporte Técnico',
    ]
    for item in toc_items:
        add_body(doc, item)

    doc.add_page_break()

    # ====== 1. INTRODUCCIÓN ======
    add_heading(doc, '1. Introducción', 1)
    add_body(doc, (
        'Bienvenido a Mercurio Software, la plataforma diseñada para facilitar el envío masivo '
        'de mensajes a través de WhatsApp Business API. Esta aplicación permite a empresas y '
        'emprendedores comunicarse de manera eficiente con sus clientes mediante mensajes '
        'personalizados, plantillas aprobadas por WhatsApp, y soporte multimedia.'
    ))
    add_body(doc, (
        'Con Mercurio Software usted podrá:'
    ))
    add_bullet(doc, 'Enviar mensajes masivos personalizados a sus clientes usando plantillas de WhatsApp')
    add_bullet(doc, 'Adjuntar documentos PDF, imágenes PNG y videos MP4 en la cabecera de sus mensajes')
    add_bullet(doc, 'Agregar hasta 4 imágenes al final de cada mensaje')
    add_bullet(doc, 'Gestionar prospectos y realizar envíos automatizados')
    add_bullet(doc, 'Mantener un historial detallado de todas las comunicaciones')
    add_bullet(doc, 'Dar soporte al cliente mediante chat en tiempo real')

    add_heading(doc, '1.1 ¿Cómo funciona?', 2)
    add_body(doc, (
        'La plataforma se conecta con la API oficial de WhatsApp Business a través de Meta '
        '(Facebook). Cada mensaje se envía utilizando plantillas previamente aprobadas por Meta, '
        'lo que garantiza que sus comunicaciones cumplan con las políticas de WhatsApp. '
        'El sistema gestiona automáticamente los límites de mensajes contratados y lleva un '
        'control detallado de cada envío.'
    ))

    doc.add_page_break()

    # ====== 2. REQUISITOS ======
    add_heading(doc, '2. Requisitos del Sistema', 1)
    add_body(doc, (
        'Para utilizar Mercurio Software, necesitará:'
    ))
    add_bullet(doc, ' Un navegador web moderno (Google Chrome, Microsoft Edge, Mozilla Firefox)')
    add_bullet(doc, ' Conexión a Internet estable')
    add_bullet(doc, ' Credenciales de acceso proporcionadas por el administrador')
    add_bullet(doc, ' Una cuenta de WhatsApp Business API configurada en Meta Developers')
    add_bullet(doc, ' Plantillas de mensajes aprobadas por Meta')

    doc.add_page_break()

    # ====== 3. INICIO DE SESIÓN ======
    add_heading(doc, '3. Inicio de Sesión', 1)
    add_body(doc, (
        'Para acceder a la plataforma, abra su navegador web y diríjase a la dirección '
        'proporcionada por su administrador. Verá la pantalla de inicio de sesión.'
    ))
    add_image(doc, '01_login.png', 'Figura 1: Pantalla de inicio de sesión')
    add_body(doc, (
        'En esta pantalla deberá ingresar:'
    ))
    add_bullet(doc, '', 'Usuario: ')
    add_body(doc, '    Su nombre de usuario asignado por el administrador.')
    add_bullet(doc, '', 'Contraseña: ')
    add_body(doc, '    La contraseña secreta proporcionada por el administrador.')

    add_body(doc, (
        'Una vez ingresados los datos, haga clic en el botón "Ingresar" para acceder al sistema. '
        'Si los datos son correctos, será redirigido a la pantalla principal. '
        'En caso de error, el sistema mostrará un mensaje indicando la causa del problema.'
    ))

    add_heading(doc, '3.1 Recuperación de Acceso', 2)
    add_body(doc, (
        'Si no tiene credenciales de acceso, haga clic en el enlace "¿No tienes cuenta? Regístrate" '
        'en la parte inferior de la pantalla de inicio de sesión. Necesitará una API Key '
        'proporcionada por el administrador para completar el registro.'
    ))

    doc.add_page_break()

    # ====== 4. PANTALLA PRINCIPAL ======
    add_heading(doc, '4. Pantalla Principal', 1)
    add_body(doc, (
        'Después de iniciar sesión, verá la pantalla principal de la aplicación. '
        'La interfaz está compuesta por los siguientes elementos:'
    ))
    add_image(doc, '02_dashboard.png', 'Figura 2: Pantalla principal (Dashboard)')

    add_heading(doc, '4.1 Barra Lateral', 2)
    add_body(doc, (
        'En el lado izquierdo de la pantalla se encuentra el menú de navegación con las '
        'siguientes opciones:'
    ))
    add_bullet(doc, ' Gestión de conversaciones de chat en tiempo real')
    add_bullet(doc, ' Creación y gestión de plantillas de WhatsApp')
    add_bullet(doc, ' Configuración de mensajes (textos e imágenes)')
    add_bullet(doc, ' Administración de prospectos y envío masivo')
    add_bullet(doc, ' Cargue de archivos multimedia')
    add_bullet(doc, ' Historial de mensajes enviados')

    add_heading(doc, '4.2 Barra Superior', 2)
    add_body(doc, (
        'En la parte superior se muestra el nombre del usuario, el plan contratado y '
        'los mensajes disponibles. También encontrará el botón "Salir" para cerrar la sesión.'
    ))

    doc.add_page_break()

    # ====== 5. GESTIÓN DE PLANTILLAS ======
    add_heading(doc, '5. Gestión de Plantillas', 1)
    add_body(doc, (
        'Las plantillas son mensajes predefinidos aprobados por Meta que puede reutilizar '
        'para comunicarse con sus clientes. Cada plantilla debe ser creada y aprobada '
        'previamente en Meta Developers antes de poder utilizarla en la plataforma.'
    ))
    add_image(doc, '06_templates.png', 'Figura 3: Pantalla de gestión de plantillas')

    add_heading(doc, '5.1 Configurar una Plantilla', 2)
    add_body(doc, (
        'Para agregar una plantilla existente a la plataforma:'
    ))
    add_bullet(doc, 'Haga clic en el botón "Agregar Plantilla"')
    add_bullet(doc, 'Complete los campos: nombre, idioma, tipo de cabecera, número de imágenes al final y número de textos variables')
    add_bullet(doc, 'Seleccione el tipo de adjunto de cabecera: Sin adjunto, Imagen, Documento o Video')
    add_bullet(doc, 'Defina cuántas imágenes aparecerán al final del mensaje (0 a 4)')
    add_bullet(doc, 'Defina cuántos textos variables tendrá el cuerpo del mensaje (0 a 6)')
    add_bullet(doc, 'Haga clic en "Guardar Plantilla"')

    add_heading(doc, '5.2 Tipos de Cabecera', 2)
    add_body(doc, (
        'Dependiendo del tipo de plantilla creada en Meta, puede seleccionar:'
    ))
    add_bullet(doc, ' — Solo texto, sin archivos adjuntos')
    add_bullet(doc, ' — Permite adjuntar una imagen PNG')
    add_bullet(doc, ' — Permite adjuntar un documento PDF')
    add_bullet(doc, ' — Permite adjuntar un video MP4')

    doc.add_page_break()

    # ====== 6. CONFIGURACIÓN DE MENSAJES ======
    add_heading(doc, '6. Configuración de Mensajes', 1)
    add_body(doc, (
        'En esta pantalla puede definir los valores por defecto para los mensajes que '
        'enviará a sus prospectos. Estos valores se aplicarán automáticamente a todos '
        'los prospectos que no tengan un valor específico configurado.'
    ))
    add_image(doc, '05_send.png', 'Figura 4: Configuración de mensajes')

    add_heading(doc, '6.1 Campos Configurables', 2)
    add_body(doc, (
        'Para cada plantilla seleccionada, podrá configurar:'
    ))
    add_bullet(doc, ' — Nombre del archivo adjunto que aparecerá en la cabecera del mensaje')
    add_bullet(doc, ' — Textos variables que se insertarán en el cuerpo del mensaje (Texto 1 a Texto 6)')
    add_bullet(doc, ' — Nombres de archivo de las imágenes que aparecerán al final del mensaje')
    add_bullet(doc, ' — Texto opcional que acompañará a cada imagen final')

    add_body(doc, (
        'Importante: Los nombres de archivo deben corresponder a archivos previamente '
        'cargados en la opción "Cargue de Archivos" del menú lateral.'
    ))

    doc.add_page_break()

    # ====== 7. PROSPECTOS ======
    add_heading(doc, '7. Prospectos y Envío Masivo', 1)
    add_body(doc, (
        'La pantalla de prospectos le permite administrar los destinatarios de sus mensajes '
        'y realizar envíos masivos de forma automatizada.'
    ))
    add_image(doc, '04_prospects.png', 'Figura 5: Pantalla de prospectos')

    add_heading(doc, '7.1 Agregar un Prospecto', 2)
    add_body(doc, (
        'Haga clic en el botón "Agregar" para crear un nuevo prospecto. Complete los campos:'
    ))
    add_bullet(doc, ' — Nombre del destinatario')
    add_bullet(doc, ' — Número de teléfono (ej. 573001234567)')
    add_bullet(doc, ' — Nombre del archivo adjunto de cabecera (si aplica)')
    add_bullet(doc, ' — Nombres de archivo de imágenes al final del mensaje')
    add_bullet(doc, ' — Textos personalizados (opcional, si no se especifica se usará el valor por defecto)')

    add_heading(doc, '7.2 Importar Prospectos desde CSV', 2)
    add_body(doc, (
        'Puede importar múltiples prospectos desde un archivo CSV:'
    ))
    add_bullet(doc, 'Haga clic en "Importar" y seleccione su archivo CSV')
    add_bullet(doc, 'El sistema validará los datos y mostrará el progreso de la importación')
    add_bullet(doc, 'Puede descargar una plantilla CSV de ejemplo con el botón "Descargar plantilla CSV"')

    add_heading(doc, '7.3 Enviar Mensajes', 2)
    add_body(doc, (
        'Para realizar un envío masivo:'
    ))
    add_bullet(doc, 'Seleccione la plantilla a utilizar en el campo "Plantilla para enviar"')
    add_bullet(doc, 'Configure los valores por defecto en "Configuración de Mensajes" si es necesario')
    add_bullet(doc, 'Seleccione "Solo pendientes" para enviar solo a prospectos no procesados')
    add_bullet(doc, 'Haga clic en "Enviar mensajes"')
    add_bullet(doc, 'El sistema mostrará el progreso del envío en tiempo real')

    add_body(doc, (
        'El flujo de envío consiste en:'
    ))
    add_bullet(doc, 'Primero se envía la plantilla con el adjunto de cabecera y textos configurados')
    add_bullet(doc, 'Luego, se envían las imágenes finales como mensajes separados (con 3 segundos de intervalo)')
    add_bullet(doc, 'Cada imagen puede incluir un texto descriptivo (caption)')

    doc.add_page_break()

    # ====== 8. CHAT ======
    add_heading(doc, '8. Chat en Tiempo Real', 1)
    add_body(doc, (
        'El módulo de chat le permite ver y responder conversaciones con sus clientes '
        'en tiempo real. Todas las interacciones (mensajes entrantes y salientes) se '
        'registran automáticamente.'
    ))
    add_image(doc, '03_chat.png', 'Figura 6: Pantalla de chat')

    add_heading(doc, '8.1 Funcionalidades del Chat', 2)
    add_bullet(doc, 'Visualización de conversaciones activas y cerradas')
    add_bullet(doc, 'Indicadores de mensajes no leídos')
    add_bullet(doc, 'Reasignación de conversaciones a otros usuarios')
    add_bullet(doc, 'Envío de mensajes de texto y archivos multimedia')
    add_bullet(doc, 'Actualización automática cada 10 segundos')

    doc.add_page_break()

    # ====== 9. CARGUE DE ARCHIVOS ======
    add_heading(doc, '9. Cargue de Archivos', 1)
    add_body(doc, (
        'Esta opción le permite subir los archivos que utilizará como adjuntos en sus '
        'mensajes. Los archivos se almacenan en la nube y se referencian por su nombre '
        'en las pantallas de Configuración y Prospectos.'
    ))
    add_image(doc, '07_upload.png', 'Figura 7: Pantalla de cargue de archivos')

    add_heading(doc, '9.1 Formatos Aceptados', 2)
    add_bullet(doc, ' — Para imágenes de cabecera e imágenes finales')
    add_bullet(doc, ' — Para documentos adjuntos en la cabecera')
    add_bullet(doc, ' — Para videos adjuntos en la cabecera')

    add_heading(doc, '9.2 Cómo Cargar Archivos', 2)
    add_bullet(doc, 'Arrastre los archivos a la zona punteada o haga clic para seleccionarlos')
    add_bullet(doc, 'Puede seleccionar múltiples archivos a la vez')
    add_bullet(doc, 'El sistema validará el tipo de archivo antes de subirlo')
    add_bullet(doc, 'Una vez subidos, verá el nombre del archivo en la tabla de resultados')
    add_bullet(doc, 'Use ese nombre exacto en las pantallas de Configuración y Prospectos')

    add_body(doc, (
        'Nota: Si sube un archivo con el mismo nombre de uno existente, el sistema '
        'lo reemplazará automáticamente.'
    ))

    doc.add_page_break()

    # ====== 10. HISTORIAL ======
    add_heading(doc, '10. Historial de Mensajes', 1)
    add_body(doc, (
        'El historial le permite consultar todos los mensajes enviados a través de la '
        'plataforma. Puede filtrar por cliente, fechas y estados.'
    ))
    add_image(doc, '08_history.png', 'Figura 8: Pantalla de historial')

    add_heading(doc, '10.1 Tipos de Historial', 2)
    add_bullet(doc, ' — Muestra los mensajes enviados con detalles de entrega')
    add_bullet(doc, ' — Vista agregada por estado de mensaje (solo administradores)')
    add_bullet(doc, ' — Listado detallado con información del cliente (solo administradores)')

    doc.add_page_break()

    # ====== 11. SOPORTE ======
    add_heading(doc, '11. Soporte Técnico', 1)
    add_body(doc, (
        'Si experimenta problemas técnicos o tiene preguntas sobre el funcionamiento '
        'de la plataforma, contacte al administrador del sistema.'
    ))
    add_body(doc, (
        'Errores comunes y soluciones:'
    ))

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Problema', 'Solución']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    rows_data = [
        ('No puedo iniciar sesión', 'Verifique usuario y contraseña. Si olvidó su contraseña, contacte al administrador.'),
        ('El mensaje no se envía', 'Verifique que tenga mensajes disponibles en su plan. Revise que la plantilla esté aprobada por Meta.'),
        ('Error al subir archivo', 'Verifique que el archivo sea PNG, PDF o MP4. El tamaño máximo es 100MB.'),
        ('La imagen no llega al cliente', 'Verifique que el nombre del archivo sea exactamente el mismo que aparece después de subirlo en "Cargue de Archivos".'),
    ]
    for row_idx, (problem, solution) in enumerate(rows_data):
        table.rows[row_idx + 1].cells[0].text = problem
        table.rows[row_idx + 1].cells[1].text = solution

    doc.add_paragraph()

    add_heading(doc, '11.1 Glosario de Términos', 2)
    glossary = [
        ('API Key', 'Clave de acceso para conectar la aplicación con servicios externos.'),
        ('Cabecera', 'Archivo multimedia (imagen, documento o video) que aparece en la parte superior del mensaje.'),
        ('Plantilla', 'Formato de mensaje predefinido y aprobado por Meta para su uso en WhatsApp Business.'),
        ('Prospecto', 'Destinatario potencial de sus mensajes.'),
        ('Footer', 'Imágenes que aparecen al final del mensaje, después del texto principal.'),
        ('Caption', 'Texto descriptivo que acompaña a una imagen en un mensaje.'),
    ]
    for term, defn in glossary:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(term + ': ')
        run.bold = True
        p.add_run(defn)

    # Set column widths for the error table
    for row in table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(12)

    # Save
    doc.save(OUTPUT_PATH)
    print(f'Documento guardado en: {OUTPUT_PATH}')

if __name__ == '__main__':
    build_document()
